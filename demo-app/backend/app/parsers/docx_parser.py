from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Optional, Union

log = logging.getLogger(__name__)

DocxSource = Union[bytes, io.BufferedReader, io.BytesIO, str]


@dataclass
class ParsedDocx:
    text: str
    paragraphs: list[str] = field(default_factory=list)
    headings: list[tuple[int, str]] = field(default_factory=list)


def parse_docx(source: DocxSource) -> ParsedDocx:
    """Extract paragraphs and headings from a DOCX file.

    Headings (Heading 1/2/3...) are preserved so the clause splitter can use
    them as strong section boundaries. Paragraphs that aren't styled as
    Headings but look like inline Title-Case section labels (e.g.
    "Termination. This Contract may be terminated...") are detected and
    promoted to numbered headings — many real contracts ship without Word
    Heading styles, and without this lift the whole body collapses into
    a single mega-clause. Tables are flattened row-by-row into the paragraph
    stream so their text still gets analyzed.
    """
    try:
        from docx import Document  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required to parse DOCX files. Install with: pip install python-docx"
        ) from exc

    handle = _to_file_like(source)
    doc = Document(handle)

    paragraphs: list[str] = []
    headings: list[tuple[int, str]] = []
    inline_heading_count = 0

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = (para.style.name or "") if para.style else ""
        level = _heading_level(style_name)
        if level is not None:
            headings.append((level, text))
            paragraphs.append(f"{'#' * min(level, 6)} {text}")
            continue

        heading, body = _split_inline_heading(text)
        if heading is not None:
            inline_heading_count += 1
            headings.append((2, heading))
            paragraphs.append(f"{inline_heading_count}. {heading}")
            if body:
                paragraphs.append(body)
            continue

        paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                paragraphs.append(row_text)

    return ParsedDocx(
        text="\n\n".join(paragraphs),
        paragraphs=paragraphs,
        headings=headings,
    )


# Connectors that can appear lowercase inside an otherwise Title-Case heading.
_HEADING_LOWERCASE_OK = frozenset(
    {
        "of", "and", "or", "the", "a", "an", "in", "on", "at", "to",
        "for", "with", "by", "as", "but", "nor", "from", "into", "upon",
    }
)

# "Heading. body..." or "Heading." — captures the candidate heading (no period)
# and any inline body that follows.
_INLINE_HEADING_RE = re.compile(
    r"^([A-Z][^\n.]{0,79})\.(?:\s+(.*))?$",
    re.DOTALL,
)


def _split_inline_heading(paragraph: str) -> tuple[Optional[str], Optional[str]]:
    """Detect a 'Title Case Heading. [body]' paragraph.

    Returns (heading, body) when the paragraph matches; (None, None) otherwise.
    `body` is None when the heading stood alone in its paragraph.
    """
    m = _INLINE_HEADING_RE.match(paragraph)
    if not m:
        return None, None
    candidate = m.group(1).strip()
    rest = (m.group(2) or "").strip() or None
    if not _looks_like_title_case_heading(candidate):
        return None, None
    return candidate, rest


def _looks_like_title_case_heading(text: str) -> bool:
    """Title-Case-y phrase, short enough to plausibly be a section label."""
    if len(text) > 80:
        return False
    words = text.split()
    if not words or len(words) > 10:
        return False
    # One-word "headings" must be substantial — keeps "Mr.", "Dr.", "Inc."
    # out without false-flagging "Confidentiality", "Termination", etc.
    if len(words) == 1 and len(text) < 7:
        return False
    for w in words:
        if w[:1].isupper():
            continue
        if w.lower().strip(".,'\"") in _HEADING_LOWERCASE_OK:
            continue
        # A lowercase non-connector word means this is prose, not a heading.
        return False
    return True


def _heading_level(style_name: str) -> int | None:
    name = style_name.strip().lower()
    if name.startswith("heading"):
        tail = name.replace("heading", "").strip()
        if tail.isdigit():
            return int(tail)
        return 1
    if name == "title":
        return 1
    return None


def _to_file_like(source: DocxSource):
    if isinstance(source, bytes):
        return io.BytesIO(source)
    return source
